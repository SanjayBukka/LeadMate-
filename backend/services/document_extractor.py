"""
Document Text Extraction Service
Extracts text from various document formats (PDF, DOCX, TXT, etc.)
"""
import os
import logging
from pathlib import Path
from typing import Optional
import PyPDF2
import docx

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text content from various document formats"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean text by removing problematic Unicode characters"""
        if not text:
            return ""
        
        # Remove or replace problematic characters
        # Replace surrogate pairs and other problematic Unicode
        cleaned = text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
        
        # Remove null bytes and other control characters except newlines and tabs
        cleaned = ''.join(char for char in cleaned if char == '\n' or char == '\t' or ord(char) >= 32)
        
        return cleaned
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file with robust error handling"""
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"Processing PDF with {total_pages} pages...")
                
                for page_num in range(total_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        
                        # Clean the extracted text to remove problematic characters
                        cleaned_text = DocumentExtractor.clean_text(text)
                        
                        if cleaned_text and cleaned_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{cleaned_text}")
                            logger.debug(f"Extracted {len(cleaned_text)} characters from page {page_num + 1}")
                    
                    except Exception as page_error:
                        logger.warning(f"Error extracting page {page_num + 1}: {page_error}")
                        text_content.append(f"--- Page {page_num + 1} ---\n[Error extracting this page]")
                        continue
                
            if not text_content:
                logger.warning("No text content extracted from PDF")
                return "[No readable text content found in PDF]"
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(full_text)} characters from PDF with {total_pages} pages")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}", exc_info=True)
            return f"[Error extracting PDF content: {str(e)}]"
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        paragraphs.append(row_text)
            
            full_text = "\n\n".join(paragraphs)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return f"[Error extracting DOCX content: {str(e)}]"
    
    @staticmethod
    def extract_from_text(file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                        logger.info(f"Extracted {len(text)} characters from text file using {encoding}")
                        return text
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail
            logger.warning("Could not decode text file with standard encodings")
            return "[Could not decode text file content]"
            
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return f"[Error extracting text content: {str(e)}]"
    
    @staticmethod
    def extract_text(file_path: str, content_type: str) -> Optional[str]:
        """
        Extract text from document based on content type
        
        Args:
            file_path: Path to the document file
            content_type: MIME type of the document
            
        Returns:
            Extracted text content or None if unsupported format
        """
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        file_extension = file_path_obj.suffix.lower()
        
        try:
            # PDF files
            if 'pdf' in content_type.lower() or file_extension == '.pdf':
                return DocumentExtractor.extract_from_pdf(file_path)
            
            # Word documents
            elif ('word' in content_type.lower() or 
                  'document' in content_type.lower() or 
                  file_extension in ['.docx', '.doc']):
                return DocumentExtractor.extract_from_docx(file_path)
            
            # Plain text files
            elif ('text' in content_type.lower() or 
                  file_extension in ['.txt', '.md', '.markdown', '.log']):
                return DocumentExtractor.extract_from_text(file_path)
            
            # Unsupported format
            else:
                logger.info(f"Unsupported file type: {content_type} ({file_extension})")
                return f"[{file_extension.upper()} files are not supported for text extraction]"
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return f"[Error extracting content: {str(e)}]"


# Global instance
document_extractor = DocumentExtractor()

